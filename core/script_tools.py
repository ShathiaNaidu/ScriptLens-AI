from __future__ import annotations

import io
import re
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from pypdf import PdfReader


SCENE_PREFIXES = ("INT.", "EXT.", "INT/EXT.", "EXT/INT.", "I/E.", "E/I.")
TRANSITIONS = ("CUT TO:", "FADE OUT.", "FADE TO BLACK.", "DISSOLVE TO:", "SMASH CUT TO:")


def read_script_file(name: str, data: bytes) -> str:
    suffix = Path(name).suffix.lower()
    if suffix in {".txt", ".fountain", ".md"}:
        return data.decode("utf-8", errors="replace")
    if suffix == ".docx":
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    if suffix == ".fdx":
        root = ET.fromstring(data)
        lines: list[str] = []
        for paragraph in root.iter("Paragraph"):
            text = "".join(node.text or "" for node in paragraph.iter("Text")).strip()
            if text:
                lines.append(text)
                lines.append("")
        return "\n".join(lines).strip()
    raise ValueError("Unsupported script file. Use PDF, DOCX, TXT, Fountain, FDX, or Markdown.")


def classify_line(line: str, previous_nonblank: str = "") -> str:
    stripped = line.strip()
    upper = stripped.upper()
    if not stripped:
        return "blank"
    if upper.startswith(SCENE_PREFIXES):
        return "scene"
    if upper in TRANSITIONS or upper.endswith(" TO:"):
        return "transition"
    if stripped.startswith("(") and stripped.endswith(")"):
        return "parenthetical"
    if len(stripped) <= 35 and stripped == upper and re.search(r"[A-Z]", stripped) and not stripped.endswith((".", "!", "?")):
        return "character"
    if previous_nonblank and classify_line(previous_nonblank) in {"character", "parenthetical"}:
        return "dialogue"
    return "action"


def normalize_screenplay(text: str) -> str:
    """Create clean Fountain-friendly screenplay text without inventing content."""
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    output: list[str] = []
    previous_nonblank = ""
    for line in lines:
        stripped = line.strip()
        kind = classify_line(stripped, previous_nonblank)
        if kind == "blank":
            if output and output[-1] != "":
                output.append("")
            continue
        if kind in {"scene", "transition", "character"}:
            stripped = stripped.upper()
        output.append(stripped)
        previous_nonblank = stripped
    return "\n".join(output).strip() + "\n"


def _iter_classified(text: str):
    previous_nonblank = ""
    for line in normalize_screenplay(text).splitlines():
        kind = classify_line(line, previous_nonblank)
        yield kind, line
        if line.strip():
            previous_nonblank = line


def export_docx(text: str, title: str = "Cinevora Screenplay") -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)
    doc.core_properties.title = title

    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(12)

    for kind, line in _iter_classified(text):
        paragraph = doc.add_paragraph()
        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        if kind == "blank":
            paragraph.add_run("")
            continue
        if kind == "scene":
            pf.space_before = Pt(12)
            run = paragraph.add_run(line.upper())
            run.bold = True
        elif kind == "character":
            pf.left_indent = Inches(2.2)
            pf.space_before = Pt(8)
            paragraph.add_run(line.upper())
        elif kind == "parenthetical":
            pf.left_indent = Inches(1.6)
            pf.right_indent = Inches(1.6)
            paragraph.add_run(line)
        elif kind == "dialogue":
            pf.left_indent = Inches(1.0)
            pf.right_indent = Inches(1.4)
            paragraph.add_run(line)
        elif kind == "transition":
            pf.alignment = 2
            paragraph.add_run(line.upper())
        else:
            paragraph.add_run(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_pdf(text: str, title: str = "Cinevora Screenplay") -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    pdf.setTitle(title)
    font = "Courier"
    size = 12
    line_height = 14
    y = height - 72
    pdf.setFont(font, size)

    def new_page():
        nonlocal y
        pdf.showPage()
        pdf.setFont(font, size)
        y = height - 72

    for kind, line in _iter_classified(text):
        if y < 72:
            new_page()
        if kind == "blank":
            y -= line_height
            continue
        if kind == "character":
            x = 270
            max_chars = 28
        elif kind == "parenthetical":
            x = 220
            max_chars = 36
        elif kind == "dialogue":
            x = 180
            max_chars = 46
        elif kind == "transition":
            # right-align transitions inside the screenplay text area
            x = max(90, width - 72 - stringWidth(line, font, size))
            max_chars = 40
        else:
            x = 108
            max_chars = 76
        chunks = [line[i:i + max_chars] for i in range(0, len(line), max_chars)] or [""]
        for chunk in chunks:
            if y < 72:
                new_page()
            pdf.drawString(x, y, chunk)
            y -= line_height
    pdf.save()
    return buffer.getvalue()


def export_fdx(text: str, title: str = "Cinevora Screenplay") -> bytes:
    root = ET.Element("FinalDraft", {"DocumentType": "Script", "Template": "No", "Version": "5"})
    content = ET.SubElement(root, "Content")
    type_map = {
        "scene": "Scene Heading",
        "action": "Action",
        "character": "Character",
        "parenthetical": "Parenthetical",
        "dialogue": "Dialogue",
        "transition": "Transition",
    }
    for kind, line in _iter_classified(text):
        if kind == "blank" or not line.strip():
            continue
        paragraph = ET.SubElement(content, "Paragraph", {"Type": type_map.get(kind, "Action")})
        text_node = ET.SubElement(paragraph, "Text")
        text_node.text = line
    tree = ET.ElementTree(root)
    out = io.BytesIO()
    tree.write(out, encoding="utf-8", xml_declaration=True)
    return out.getvalue()
